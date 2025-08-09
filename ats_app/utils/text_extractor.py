import pdfplumber
import fitz  # PyMuPDF
import docx
import docx2txt
import re
from typing import Dict, Any

class TextExtractor:
    @staticmethod
    def extract_from_pdf(file_path: str) -> Dict[str, Any]:
        """Extract text from PDF with formatting preservation"""
        text = ""
        metadata = {"pages": 0, "formatting_issues": []}
        
        try:
            # Try pdfplumber first (better for structured text)
            with pdfplumber.open(file_path) as pdf:
                metadata["pages"] = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    else:
                        metadata["formatting_issues"].append(f"Page {page.page_number} has extraction issues")
        except Exception as e:
            # Fallback to PyMuPDF
            try:
                doc = fitz.open(file_path)
                metadata["pages"] = doc.page_count
                for page_num in range(doc.page_count):
                    page = doc[page_num]
                    text += page.get_text() + "\n"
                doc.close()
            except Exception as e2:
                raise Exception(f"Failed to extract PDF: {str(e)}, {str(e2)}")
        
        return {
            "text": text.strip(),
            "metadata": metadata
        }
    
    @staticmethod
    def extract_from_docx(file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX with structure preservation"""
        try:
            # Extract plain text
            text = docx2txt.process(file_path)
            
            # Extract structured content for better analysis
            doc = docx.Document(file_path)
            structured_content = {
                "paragraphs": [],
                "tables": [],
                "headers": []
            }
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    structured_content["paragraphs"].append({
                        "text": paragraph.text,
                        "style": paragraph.style.name if paragraph.style else None
                    })
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                structured_content["tables"].append(table_data)
            
            metadata = {
                "paragraphs_count": len(structured_content["paragraphs"]),
                "tables_count": len(structured_content["tables"]),
                "formatting_preserved": True
            }
            
            return {
                "text": text.strip(),
                "structured_content": structured_content,
                "metadata": metadata,
                "document": doc  # Keep reference for updating
            }
            
        except Exception as e:
            raise Exception(f"Failed to extract DOCX: {str(e)}")
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep important punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)]', ' ', text)
        # Remove extra spaces
        text = ' '.join(text.split())
        return text.strip()
    
    @staticmethod
    def extract_sections(text: str) -> Dict[str, str]:
        """Extract common resume sections"""
        sections = {
            "contact": "",
            "summary": "",
            "experience": "",
            "education": "",
            "skills": "",
            "projects": "",
            "certifications": ""
        }
        
        # Define section patterns
        section_patterns = {
            "contact": r"(contact|phone|email|address).*?(?=\n\n|\n[A-Z])",
            "summary": r"(summary|objective|profile).*?(?=\n\n|\n[A-Z])",
            "experience": r"(experience|employment|work history).*?(?=\n\n|\n[A-Z])",
            "education": r"(education|academic).*?(?=\n\n|\n[A-Z])",
            "skills": r"(skills|technical skills|competencies).*?(?=\n\n|\n[A-Z])",
            "projects": r"(projects|portfolio).*?(?=\n\n|\n[A-Z])",
            "certifications": r"(certifications|certificates|licenses).*?(?=\n\n|\n[A-Z])"
        }
        
        text_lower = text.lower()
        for section, pattern in section_patterns.items():
            match = re.search(pattern, text_lower, re.DOTALL | re.IGNORECASE)
            if match:
                sections[section] = match.group(0).strip()
        
        return sections
    