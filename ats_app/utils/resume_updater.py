import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import openai
import os
from typing import Dict, List, Any, Optional

class ResumeUpdater:
    def __init__(self, openai_api_key: Optional[str] = None):
        if openai_api_key:
            openai.api_key = openai_api_key
            self.use_openai = True
        else:
            self.use_openai = False
    
    def update_resume(self, resume_data: Dict[str, Any], missing_skills: List[str], 
                     missing_keywords: List[str], output_path: str) -> Dict[str, Any]:
        """Update resume with missing skills and keywords"""
        if 'document' not in resume_data:
            raise ValueError("No document object found in resume data")
        
        doc = resume_data['document']
        
        # Create a copy of the document
        updated_doc = Document()
        
        # Copy styles from original document
        self._copy_styles(doc, updated_doc)
        
        # Process each paragraph
        skills_added_to_section = False
        skills_integrated_count = 0
        
        for paragraph in doc.paragraphs:
            new_paragraph = updated_doc.add_paragraph()
            
            # Copy paragraph formatting
            self._copy_paragraph_formatting(paragraph, new_paragraph)
            
            # Check if this is a skills section
            if self._is_skills_section(paragraph.text):
                # Add missing skills to skills section
                new_paragraph.text = paragraph.text
                if missing_skills and not skills_added_to_section:
                    self._add_skills_to_section(updated_doc, missing_skills)
                    skills_added_to_section = True
            
            # Check if this is an experience bullet point
            elif self._is_experience_bullet(paragraph.text):
                # Try to integrate missing skills naturally
                updated_text, integrated = self._integrate_skills_in_experience(
                    paragraph.text, missing_skills, missing_keywords
                )
                new_paragraph.text = updated_text
                if integrated:
                    skills_integrated_count += 1
            
            else:
                # Copy paragraph as-is
                new_paragraph.text = paragraph.text
        
        # If no skills section found, create one
        if not skills_added_to_section and missing_skills:
            self._create_skills_section(updated_doc, missing_skills)
        
        # Copy tables
        for table in doc.tables:
            self._copy_table(table, updated_doc)
        
        # Save updated document
        updated_doc.save(output_path)
        
        return {
            'output_path': output_path,
            'skills_added_to_section': skills_added_to_section,
            'skills_integrated_count': skills_integrated_count,
            'total_skills_added': len(missing_skills),
            'total_keywords_integrated': len(missing_keywords)
        }
    
    def _copy_styles(self, source_doc: Document, target_doc: Document):
        """Copy styles from source to target document"""
        try:
            for style in source_doc.styles:
                if style.name not in target_doc.styles:
                    # This is a simplified approach - full style copying is complex
                    pass
        except Exception:
            # If style copying fails, continue without it
            pass
    
    def _copy_paragraph_formatting(self, source_para, target_para):
        """Copy paragraph formatting"""
        try:
            # Copy basic formatting
            target_para.alignment = source_para.alignment
            
            # Copy runs with formatting
            for run in source_para.runs:
                new_run = target_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                if run.font.size:
                    new_run.font.size = run.font.size
                if run.font.name:
                    new_run.font.name = run.font.name
        except Exception:
            # If formatting copy fails, just copy text
            target_para.text = source_para.text
    
    def _copy_table(self, source_table, target_doc):
        """Copy table to target document"""
        try:
            rows = len(source_table.rows)
            cols = len(source_table.columns)
            
            table = target_doc.add_table(rows=rows, cols=cols)
            
            for i, row in enumerate(source_table.rows):
                for j, cell in enumerate(row.cells):
                    table.rows[i].cells[j].text = cell.text
        except Exception:
            # If table copying fails, skip it
            pass
    
    def _is_skills_section(self, text: str) -> bool:
        """Check if paragraph is part of skills section"""
        skills_indicators = [
            'skills', 'technical skills', 'core competencies', 'expertise',
            'technologies', 'proficiencies', 'technical proficiencies'
        ]
        
        text_lower = text.lower().strip()
        return any(indicator in text_lower for indicator in skills_indicators)
    
    def _is_experience_bullet(self, text: str) -> bool:
        """Check if paragraph is an experience bullet point"""
        text = text.strip()
        # Check for bullet point indicators
        bullet_indicators = ['•', '-', '*', '◦']
        action_words = [
            'developed', 'managed', 'led', 'created', 'implemented', 'designed',
            'achieved', 'improved', 'increased', 'reduced', 'collaborated',
            'coordinated', 'supervised', 'analyzed', 'optimized'
        ]
        
        starts_with_bullet = any(text.startswith(indicator) for indicator in bullet_indicators)
        contains_action = any(word in text.lower() for word in action_words)
        
        return starts_with_bullet or (contains_action and len(text) > 30)
    
    def _add_skills_to_section(self, doc: Document, missing_skills: List[str]):
        """Add missing skills as bullet points to skills section"""
        skills_paragraph = doc.add_paragraph()
        
        for skill in missing_skills[:10]:  # Limit to top 10 skills
            bullet_para = doc.add_paragraph(f"• {skill.title()}", style='List Bullet')
    
    def _create_skills_section(self, doc: Document, missing_skills: List[str]):
        """Create a new skills section"""
        # Add section header
        skills_header = doc.add_heading('Technical Skills', level=2)
        
        # Add skills as bullet points
        for skill in missing_skills[:12]:  # Limit to top 12 skills
            doc.add_paragraph(f"• {skill.title()}", style='List Bullet')
    
    def _integrate_skills_in_experience(self, experience_text: str, 
                                      missing_skills: List[str], 
                                      missing_keywords: List[str]) -> tuple[str, bool]:
        """Integrate missing skills naturally into experience bullet points"""
        if self.use_openai:
            return self._integrate_skills_with_openai(experience_text, missing_skills, missing_keywords)
        else:
            return self._integrate_skills_simple(experience_text, missing_skills, missing_keywords)
    
    def _integrate_skills_with_openai(self, experience_text: str, 
                                    missing_skills: List[str], 
                                    missing_keywords: List[str]) -> tuple[str, bool]:
        """Use OpenAI to naturally integrate skills"""
        try:
            skills_to_add = missing_skills[:3]  # Limit to 3 skills per bullet
            keywords_to_add = missing_keywords[:2]  # Limit to 2 keywords per bullet
            
            prompt = f"""
            Please rewrite the following experience bullet point to naturally include these missing skills and keywords while maintaining the original meaning and impact:
            
            Original bullet point: {experience_text}
            
            Skills to integrate: {', '.join(skills_to_add)}
            Keywords to integrate: {', '.join(keywords_to_add)}
            
            Guidelines:
            - Keep the original achievement and impact
            - Integrate skills naturally without forcing them
            - Maintain professional tone
            - Don't change the core message
            - Only add skills that make sense in this context
            
            Rewritten bullet point:
            """
            
            response = openai.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=150,
                temperature=0.3
            )
            
            updated_text = response.choices[0].message.content.strip()
            return updated_text, True
            
        except Exception as e:
            # Fallback to simple integration
            return self._integrate_skills_simple(experience_text, missing_skills, missing_keywords)
    
    def _integrate_skills_simple(self, experience_text: str, 
                               missing_skills: List[str], 
                               missing_keywords: List[str]) -> tuple[str, bool]:
        """Simple rule-based skill integration"""
        # Select 1-2 most relevant skills to add
        skills_to_add = missing_skills[:2]
        integrated = False
        
        updated_text = experience_text
        
        # Try to add skills in context
        for skill in skills_to_add:
            skill_lower = skill.lower()
            
            # Add programming languages
            if skill_lower in ['python', 'java', 'javascript', 'c++', 'sql']:
                if 'develop' in updated_text.lower() or 'program' in updated_text.lower():
                    updated_text = updated_text.replace(
                        'developed', f'developed using {skill}', 1
                    )
                    integrated = True
                    break
            
            # Add tools and technologies
            elif skill_lower in ['excel', 'tableau', 'powerbi', 'sql', 'git']:
                if 'analy' in updated_text.lower() or 'data' in updated_text.lower():
                    updated_text += f' utilizing {skill}'
                    integrated = True
                    break
        
        return updated_text, integrated
